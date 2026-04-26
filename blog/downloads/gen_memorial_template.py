# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
"""
Gera modelo de Memorial Descritivo de Obra editável em DOCX e PDF.
Salva em /blog/downloads/ pra ser baixado direto pelo site.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

THIS_DIR = os.path.dirname(os.path.abspath(__file__))

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x47, 0x55, 0x69)


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = INDIGO
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = INDIGO
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def add_p(doc, text, italic=False, gray=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    run.font.color.rgb = GRAY if gray else DARK
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)


def add_field(doc, label, hint=""):
    """Campo a preencher: 'LABEL: ___________'"""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.font.size = Pt(10)
    r1.bold = True
    r1.font.color.rgb = DARK
    r1.font.name = 'Calibri'
    r2 = p.add_run("[preencher]")
    r2.font.size = Pt(10)
    r2.italic = True
    r2.font.color.rgb = GRAY
    r2.font.name = 'Calibri'
    if hint:
        r3 = p.add_run(f"  ({hint})")
        r3.font.size = Pt(8)
        r3.italic = True
        r3.font.color.rgb = GRAY
        r3.font.name = 'Calibri'


def main():
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── CAPA ────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    run = p.add_run('MEMORIAL DESCRITIVO')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = INDIGO
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DE OBRA')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run('Modelo editável · Estrutura completa em 10 seções')
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'

    # Bloco de identificação
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run('Projeto: ____________________________')
    run.font.size = Pt(12)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Cliente: ____________________________')
    run.font.size = Pt(12)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Local: ______________________________')
    run.font.size = Pt(12)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Data: _______________________________')
    run.font.size = Pt(12)
    run.font.name = 'Calibri'

    # Rodapé da capa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run('Modelo gerado por AI.arq · ai.arq.br')
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'

    doc.add_page_break()

    # ── INSTRUÇÕES DE USO ──────────────────────────────
    add_h1(doc, 'Como usar este modelo')
    add_p(doc,
        'Este é um modelo genérico de memorial descritivo de obra. '
        'Foi estruturado pra atender 95% dos casos típicos do mercado brasileiro. '
        'Adapte cada seção pra realidade do seu projeto.')
    add_p(doc, 'Recomendações de preenchimento:', italic=False)
    add_p(doc, '• Substitua os textos em [colchetes] pelos dados reais do projeto.')
    add_p(doc, '• Cite normas ABNT relevantes (NBR 15575, NBR 9050, etc.) onde aplicável.')
    add_p(doc, '• Especifique sempre marca/modelo de itens críticos (louças, metais, esquadrias).')
    add_p(doc, '• Cruze a descrição com a planilha de quantitativos pra evitar inconsistência.')
    add_p(doc, '• Datilografe o documento. Memorial manuscrito não é aceito por bancos/prefeituras.')
    add_p(doc, '')
    add_p(doc, 'IMPORTANTE: Este memorial deve ser revisado e assinado por profissional habilitado (arquiteto ou engenheiro com CREA/CAU ativo). O AI.arq não substitui profissional habilitado.', italic=True, gray=True)

    doc.add_page_break()

    # ── 1. IDENTIFICAÇÃO DO PROJETO ─────────────────────
    add_h1(doc, '1. Identificação do projeto')
    add_field(doc, 'Proprietário')
    add_field(doc, 'CPF/CNPJ')
    add_field(doc, 'Endereço da obra')
    add_field(doc, 'Cidade/Estado')
    add_field(doc, 'Número do lote/Inscrição imobiliária')
    add_field(doc, 'Área do terreno', 'em m²')
    add_field(doc, 'Área construída total', 'em m²')
    add_field(doc, 'Área a construir/reformar', 'em m²')
    add_field(doc, 'Finalidade', 'residencial / comercial / industrial')
    add_field(doc, 'Profissional responsável')
    add_field(doc, 'CAU/CREA')
    add_field(doc, 'ART/RRT', 'número da anotação')

    # ── 2. CARACTERÍSTICAS GERAIS ──────────────────────
    add_h1(doc, '2. Características gerais da edificação')
    add_h3(doc, '2.1 Tipologia')
    add_p(doc, '[Descrever o tipo de edificação: térrea, sobrado, edifício multipavimentos, galpão industrial, etc. Incluir número de pavimentos, número de unidades habitacionais (se aplicável), capacidade.]')

    add_h3(doc, '2.2 Padrão construtivo')
    add_p(doc, '[Descrever padrão geral: popular, médio, alto padrão. Justificar com base nos materiais e acabamentos especificados.]')

    add_h3(doc, '2.3 Sistema construtivo')
    add_p(doc, '[Descrever o sistema: alvenaria estrutural, alvenaria com estrutura independente em concreto armado, sistema light steel frame, wood frame, etc.]')

    # ── 3. SISTEMA CONSTRUTIVO ─────────────────────────
    add_h1(doc, '3. Sistema construtivo')
    add_h3(doc, '3.1 Fundação')
    add_p(doc, '[Tipo de fundação (sapata corrida, sapata isolada, radier, estaca, tubulão), profundidade, dimensões, materiais, conforme projeto estrutural.]')

    add_h3(doc, '3.2 Estrutura')
    add_p(doc, '[Sistema estrutural (concreto armado, alvenaria estrutural, metálica, mista). Especificar fck do concreto, aço utilizado, normas seguidas (NBR 6118).]')

    add_h3(doc, '3.3 Vedação')
    add_p(doc, '[Tipo de alvenaria (cerâmica, bloco de concreto, bloco sílico-calcário). Espessuras (interna 9cm, externa 14cm). Argamassa de assentamento.]')

    add_h3(doc, '3.4 Cobertura')
    add_p(doc, '[Estrutura (madeira, metálica). Telhamento (cerâmico, fibrocimento, metálico, sandwich). Inclinação. Calhas e rufos. Impermeabilização.]')

    # ── 4. ESPECIFICAÇÃO POR AMBIENTE ─────────────────
    add_h1(doc, '4. Especificação por ambiente')
    add_p(doc, '[Para cada ambiente da obra, descrever: piso, parede, forro, esquadrias, instalações, mobiliário fixo. Repetir o bloco abaixo pra cada ambiente.]', italic=True, gray=True)

    add_h3(doc, '4.1 Sala de estar/jantar')
    add_p(doc, 'Piso: [especificar material, modelo, marca, formato, junta]')
    add_p(doc, 'Paredes: [pintura, papel de parede, revestimento — com cor e marca]')
    add_p(doc, 'Forro: [tipo, altura final, iluminação embutida]')
    add_p(doc, 'Esquadrias: [tipo P1, P2 conforme quadro]')
    add_p(doc, 'Instalações: [pontos elétricos, dados, climatização]')

    add_h3(doc, '4.2 Cozinha')
    add_p(doc, 'Piso: [especificar]')
    add_p(doc, 'Paredes: [revestimento até X metros, pintura acima]')
    add_p(doc, 'Forro: [especificar]')
    add_p(doc, 'Esquadrias: [especificar]')
    add_p(doc, 'Instalações: [pontos elétricos com tensão, pontos hidráulicos, exaustão]')
    add_p(doc, 'Bancadas: [material, espessura, comprimento]')
    add_p(doc, 'Marcenaria fixa: [especificar armários superiores e inferiores, material, ferragem]')

    add_h3(doc, '4.3 Banheiros')
    add_p(doc, 'Piso: [especificar]')
    add_p(doc, 'Paredes: [revestimento até teto, ou só área molhada]')
    add_p(doc, 'Louças: [marca, modelo de bacia, lavatório, ducha higiênica]')
    add_p(doc, 'Metais: [marca, modelo de torneiras, registros, chuveiros]')
    add_p(doc, 'Box: [tipo, vidro, ferragem]')

    add_h3(doc, '4.4 Quartos / Suítes')
    add_p(doc, 'Piso: [especificar]')
    add_p(doc, 'Paredes: [pintura, marca, cor]')
    add_p(doc, 'Forro: [especificar]')
    add_p(doc, 'Esquadrias: [janela, porta]')

    add_h3(doc, '4.5 Áreas externas')
    add_p(doc, 'Piso: [especificar — porcelanato externo, deck, concreto desempenado]')
    add_p(doc, 'Pintura: [tipo de tinta para áreas externas]')
    add_p(doc, 'Coberturas: [pergolados, marquises]')

    # ── 5. ESPECIFICAÇÃO POR DISCIPLINA ───────────────
    add_h1(doc, '5. Especificação por disciplina')
    add_h3(doc, '5.1 Piso')
    add_p(doc, '[Listar todos os tipos de piso da obra, com marca, modelo, formato, PEI (se cerâmico), espessura, junta, rejunte. Quantidade total por tipo.]')

    add_h3(doc, '5.2 Revestimento de parede')
    add_p(doc, '[Tipos de revestimento, cor, modelo, marca, área de aplicação.]')

    add_h3(doc, '5.3 Forro')
    add_p(doc, '[Tipo (gesso liso, gesso decorativo, modular, PVC, madeira), marca, espessura, altura final do pé-direito.]')

    add_h3(doc, '5.4 Pintura')
    add_p(doc, '[Tipo de tinta (látex PVA, acrílica, esmalte, epóxi), número de demãos, fundo preparador. Especificar diferente entre paredes internas, externas, teto, esquadrias.]')

    # ── 6. INSTALAÇÕES ─────────────────────────────────
    add_h1(doc, '6. Instalações')
    add_h3(doc, '6.1 Elétrica')
    add_p(doc, 'Padrão de entrada: [trifásico/bifásico/monofásico, kVA contratado]')
    add_p(doc, 'Quadro de distribuição: [marca, modelo, número de circuitos]')
    add_p(doc, 'Fiação: [seção mínima, marca]')
    add_p(doc, 'Eletrocalhas/conduítes: [tipo]')
    add_p(doc, 'Tomadas: [marca, modelo, padrão NBR 14136]')
    add_p(doc, 'Interruptores: [marca, modelo]')
    add_p(doc, 'Iluminação: [tipos de luminárias, lâmpadas LED, potências]')

    add_h3(doc, '6.2 Hidráulica')
    add_p(doc, 'Tubulação água fria: [PVC, PEX, PPR, marca, diâmetros]')
    add_p(doc, 'Tubulação água quente: [CPVC, PEX, PPR]')
    add_p(doc, 'Reservatórios: [capacidade total, posição]')
    add_p(doc, 'Aquecimento: [boiler elétrico, gás, solar]')

    add_h3(doc, '6.3 Esgoto')
    add_p(doc, 'Tubulação: [PVC série normal, série esgoto, diâmetros]')
    add_p(doc, 'Caixa de gordura: [especificar]')
    add_p(doc, 'Sumidouro/fossa séptica: [se aplicável]')
    add_p(doc, 'Caixas de inspeção: [posição, dimensões]')

    add_h3(doc, '6.4 Climatização (HVAC)')
    add_p(doc, 'Sistema: [splits, VRF, central]')
    add_p(doc, 'Capacidade total: [BTUs ou TR]')
    add_p(doc, 'Marca/modelo: [especificar]')

    add_h3(doc, '6.5 Combate a incêndio')
    add_p(doc, '[Sprinkler, hidrantes, extintores, sinalização — conforme exigência do corpo de bombeiros local.]')

    # ── 7. ESQUADRIAS ──────────────────────────────────
    add_h1(doc, '7. Esquadrias')
    add_p(doc, 'Conforme quadro de esquadrias na prancha [A-XXX]:')
    add_p(doc, 'Portas internas: [material — madeira maciça/lisa, MDF; abertura; ferragem; acabamento]')
    add_p(doc, 'Portas externas: [material — madeira maciça, alumínio, aço; ferragem; pintura/anodização]')
    add_p(doc, 'Janelas: [material — alumínio, madeira, PVC; tipo de abertura; vidro — comum/laminado/temperado/duplo; espessura]')
    add_p(doc, 'Esquadrias especiais: [persianas, brises, divisórias]')

    # ── 8. PINTURA ─────────────────────────────────────
    add_h1(doc, '8. Pintura')
    add_p(doc, 'Paredes internas: [tinta látex PVA, marca, número de demãos, fundo preparador]')
    add_p(doc, 'Paredes externas: [tinta acrílica resistente a intempéries]')
    add_p(doc, 'Teto: [tinta látex branca]')
    add_p(doc, 'Esquadrias de madeira: [esmalte sintético, número de demãos, primer]')
    add_p(doc, 'Esquadrias metálicas: [esmalte sintético, primer anti-corrosivo]')
    add_p(doc, 'Cores: [definidas no projeto de interiores]')

    # ── 9. ÁREAS EXTERNAS / PAISAGISMO ─────────────────
    add_h1(doc, '9. Áreas externas')
    add_p(doc, 'Pavimentação: [especificar tipo e local]')
    add_p(doc, 'Drenagem: [calhas, ralos, captação de água pluvial]')
    add_p(doc, 'Paisagismo: [conforme projeto específico]')
    add_p(doc, 'Cercamento: [muros, gradis, portões]')

    # ── 10. CONSIDERAÇÕES FINAIS ───────────────────────
    add_h1(doc, '10. Considerações finais')
    add_p(doc, 'Este memorial descritivo é parte integrante do projeto arquitetônico e deve ser lido em conjunto com:')
    add_p(doc, '• Pranchas arquitetônicas (plantas, cortes, fachadas, detalhes)')
    add_p(doc, '• Planilha de quantitativos')
    add_p(doc, '• Projeto estrutural')
    add_p(doc, '• Projeto elétrico, hidráulico e demais complementares')
    add_p(doc, '• ART/RRT do responsável técnico')
    add_p(doc, '')
    add_p(doc, 'Qualquer alteração na execução deve ser previamente aprovada pelo profissional responsável.')
    add_p(doc, '')
    add_p(doc, 'Todas as obras devem seguir as normas técnicas brasileiras vigentes (ABNT NBR 15575 — Desempenho, NBR 9050 — Acessibilidade, NBR 6118 — Concreto armado, NBR 5410 — Instalações elétricas de baixa tensão, e demais aplicáveis).')

    # Assinaturas
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('___________________________________')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Profissional Responsável\nCAU/CREA: __________')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Local e data: ___________________________, ____ de ____________ de 20____')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    # Rodapé final
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Este modelo foi gerado por AI.arq.')
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Pra acelerar a planilha de quantitativos, acesse ai.arq.br')
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = INDIGO
    run.font.name = 'Calibri'

    docx_path = os.path.join(THIS_DIR, "memorial-descritivo-obra-modelo.docx")
    doc.save(docx_path)
    print(f"✓ DOCX salvo: {docx_path}")


if __name__ == "__main__":
    main()
