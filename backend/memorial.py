# -*- coding: utf-8 -*-
"""Memorial descritivo (RASCUNHO) — gerado a partir do quantitativo do projeto.

v1 DETERMINÍSTICA (01/08/2026, estudo em docs/ESTUDO_MEMORIAL_DESCRITIVO_2026-08.md):
nenhuma chamada de IA — o texto é template + os PRÓPRIOS itens do projeto.
Zero risco de inventar especificação, marca ou norma (as 4 cercas do estudo):

  1. Medido × estimado ROTULADO no texto (regra dura nº1 aplicada a prosa);
  2. O que o CAD não tem vira [A PREENCHER PELO RESPONSÁVEL TÉCNICO] explícito;
  3. Zero invenção: nenhuma norma/marca citada que não venha dos itens;
  4. Carimbo RASCUNHO em cabeçalho, capa e rodapé — não substitui o memorial
     assinado (RRT/ART). Regra dura nº5: não substitui o profissional.

v1.1 (mesmo dia): separado em DUAS etapas pra tela editável (memorial.html):
  montar_estrutura(projeto, items) -> dict JSON (o que a tela edita e o banco salva)
  estrutura_para_docx(caminho, estrutura) -> renderiza o Word a partir do JSON
gerar_memorial_docx() (compat) = montar + renderizar.

⚠️ NBR 13531/13532 estão CANCELADAS (dez/2017 → série NBR 16636). Este módulo
não cita norma nenhuma de propósito — só o que estiver escrito nos itens.
"""

import re
import re as _re   # 🪤 21/08: duas funções usavam `re` sem alias — NameError escondido
from datetime import datetime, timezone, timedelta

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ordem construtiva das disciplinas (valores REAIS do banco, conferidos 01/08).
# Disciplina fora da lista entra no fim, antes de "Complementares".
ORDEM_DISCIPLINAS = [
    "Demolição e Remoção",
    "Serviços Preliminares",
    "Estrutura",
    "Fechamentos Verticais",
    "Divisórias e Vidros",
    "Portas e Ferragens",
    "Instalações Elétricas e Dados",
    "Iluminação",
    "Instalações Hidráulicas",
    "Instalações de Gás",
    "Ar-Condicionado",
    "Incêndio e Segurança",
    "Revestimentos",
    "Forros",
    "Pisos e Rodapés",
    "Marcenaria",
    "Mobiliário",
    "Persianas e Cortinas",
    "Complementares",
]

# Lacunas específicas por disciplina — o que um memorial de verdade tem e o
# CAD NÃO tem. Sempre em bloco [A PREENCHER] explícito (cerca nº2).
LACUNAS_POR_DISCIPLINA = {
    "Estrutura": ("características do concreto (fck e traço), resultado da sondagem do solo "
                  "e dimensionamento estrutural conforme projeto específico"),
    "Fechamentos Verticais": ("traço da argamassa de assentamento e de revestimento, e "
                              "especificação do bloco/tijolo quando não indicada em projeto"),
    "Revestimentos": "preparo de base, argamassa colante e rejunte",
    "Pisos e Rodapés": "preparo de contrapiso, argamassa colante e rejunte",
    "Instalações Hidráulicas": "pressões de teste, caimentos e detalhes executivos",
    "Instalações Elétricas e Dados": "quadros de cargas, balanceamento e aterramento",
}

LACUNA_PADRAO = "especificações complementares, método executivo e marcas/modelos de referência"

MARCA_PREENCHER = "[A PREENCHER PELO RESPONSÁVEL TÉCNICO"

CINZA = RGBColor(0x6B, 0x72, 0x80)
LARANJA = RGBColor(0xB4, 0x5B, 0x09)
VERMELHO = RGBColor(0xB9, 0x1C, 0x1C)


def _agora_brasilia():
    # Horários sempre em Brasília (regra da casa) — UTC-3 sem depender de tz do host.
    return datetime.now(timezone.utc) - timedelta(hours=3)


def _fmt_qty(q):
    """1234.5 → '1.234,5' (pt-BR, até 2 casas, sem zeros à direita)."""
    try:
        q = float(q or 0)
    except (TypeError, ValueError):
        return "0"
    txt = f"{q:,.2f}".replace(",", "§").replace(".", ",").replace("§", ".")
    if "," in txt:
        txt = txt.rstrip("0").rstrip(",")
    return txt or "0"


# ═══════════════════════════════════════════════════════════════
#  ETAPA 1 — estrutura JSON (o que a tela edita e o banco salva)
# ═══════════════════════════════════════════════════════════════

def _bloco_preencher(detalhe):
    return {"tipo": "preencher", "texto": f"{MARCA_PREENCHER}: {detalhe}.]"}


_RX_SEG_ABERTO = _re.compile(
    r"^(?:material|tipo|cor|dimens[aã]o|especifica[cç][aã]o|fabricante|modelo|capacidade|acabamento)?"
    r"[^,;.]{0,40}?a (?:definir|especificar|confirmar)(?=$|[\s,;.)])", _re.I)


def _poda_boilerplate(desc: str) -> str:
    """Remove da PROSA do memorial os segmentos "— material a especificar —"
    que a IA repete em quase todo item (20/08/2026: no projeto de teste, o eco
    aparecia em ~todas as linhas e o documento inteiro gritava rascunho).

    🪤 Só na renderização do memorial — a planilha continua intacta. E só cai o
    SEGMENTO entre travessões que é puro "a definir/especificar"; segmento com
    conteúdo real ("vidro temperado 10 mm") fica. A lacuna continua declarada
    UMA vez por seção, no bloco [A PREENCHER] que já existe.
    """
    partes = [p.strip() for p in _re.split(r"\s+—\s+", desc or "") if p.strip()]
    if len(partes) <= 1:
        return (desc or "").strip()
    mantidas = [partes[0]] + [p for p in partes[1:] if not _RX_SEG_ABERTO.match(p)]
    return " — ".join(mantidas)


def _agrupar_por_prefixo(itens: list, minimo: int = 4):
    """Agrupa itens que compartilham prefixo (antes do 1º "—") e unidade.

    20/08/2026: a seção de Pisos do memorial tinha 24 blocos, um por ambiente
    — lista telefônica, não prosa. Com >=`minimo` itens do mesmo prefixo+unidade,
    viram UM bloco com total e ambientes listados.
    Devolve lista de ("grupo", prefixo_exibicao, itens) | ("solto", None, [item]),
    na ordem de chegada.
    """
    buckets, ordem = {}, []
    for it in itens:
        d = (it.get("description") or "").strip()
        pref = _re.split(r"\s+—\s+", d)[0].strip()
        k = (pref.lower(), (it.get("unit") or "").strip().lower())
        if k not in buckets:
            buckets[k] = []
            ordem.append((k, pref))
        buckets[k].append(it)
    saida = []
    for k, pref in ordem:
        grupo = buckets[k]
        if len(grupo) >= minimo and pref:
            saida.append(("grupo", pref, grupo))
        else:
            for it in grupo:
                saida.append(("solto", None, [it]))
    return saida


def montar_estrutura(projeto: dict, items: list) -> dict:
    """Monta a estrutura editável do memorial a partir do projeto + itens.

    Tipos de bloco: texto · preencher · item (origem: medido|estimado) ·
    tabela (linhas [rótulo, valor]) · assinatura. A tela edita `texto`/`linhas`;
    títulos e ordem ficam fixos no v1.
    """
    nome_obra = (projeto.get("project_name") or "").strip() or "[A PREENCHER: nome da obra]"
    total = len(items)
    medidos = sum(1 for i in items if (i.get("confidence") or "") == "confirmado")
    estimados = total - medidos

    secoes = []
    secoes.append({"titulo": "1. Apresentação", "blocos": [
        {"tipo": "texto", "texto": (
            "Este memorial descritivo relaciona os serviços e materiais identificados no projeto, "
            "a partir da leitura do arquivo CAD (plantas e pranchas) enviado ao AI.arq. "
            "Ele foi organizado na sequência usual das etapas construtivas e serve como BASE para o "
            "memorial definitivo, a ser complementado, corrigido e assinado pelo responsável técnico.")},
        {"tipo": "texto", "texto": (
            f"O levantamento contém {total} itens: {medidos} com quantidade medida diretamente da "
            f"geometria do CAD e {estimados} com quantidade estimada ou a confirmar. No texto, cada "
            "item indica sua origem entre parênteses — “medido do CAD” ou “estimativa — a confirmar”. "
            "Nenhuma especificação, marca ou norma foi acrescentada além do que consta no próprio projeto; "
            "onde o memorial exige informação que não está no CAD, há um campo destacado "
            "[A PREENCHER PELO RESPONSÁVEL TÉCNICO].")},
    ]})

    area_txt = None
    if projeto.get("total_area"):
        area_txt = f"{_fmt_qty(projeto['total_area'])} m² (medida no projeto)"
    elif projeto.get("user_total_area"):
        area_txt = f"{_fmt_qty(projeto['user_total_area'])} m² (informada pelo cliente — não medida)"
    secoes.append({"titulo": "2. Dados da obra", "blocos": [
        {"tipo": "tabela", "linhas": [
            ["Obra / projeto", nome_obra],
            ["Tipologia", (projeto.get("typology") or "").strip() or "[A PREENCHER]"],
            ["Área total", area_txt or "[A PREENCHER]"],
            ["Endereço", "[A PREENCHER]"],
            ["Proprietário / contratante", "[A PREENCHER]"],
        ]},
    ]})

    secoes.append({"titulo": "3. Responsabilidade técnica", "blocos": [
        _bloco_preencher("nome do responsável técnico, título profissional, registro CAU/CREA e número da RRT/ART"),
    ]})

    # Seções técnicas por disciplina, em ordem construtiva
    grupos = {}
    for it in items:
        d = (it.get("discipline") or "").strip() or "Outros itens levantados"
        grupos.setdefault(d, []).append(it)
    conhecidas = [d for d in ORDEM_DISCIPLINAS if d in grupos]
    extras = sorted(d for d in grupos if d not in ORDEM_DISCIPLINAS)
    if "Complementares" in conhecidas:
        idx = conhecidas.index("Complementares")
        ordem = conhecidas[:idx] + extras + conhecidas[idx:]
    else:
        ordem = conhecidas + extras

    num = 4
    for disc in ordem:
        blocos = [{"tipo": "texto", "texto": (
            f"Os serviços de {disc.lower()} compreendem os itens identificados no projeto, "
            "conforme relação abaixo:")}]
        _ordenados = sorted(grupos[disc], key=lambda x: (x.get("sort_order") or 0, x.get("item_num") or ""))

        def _qtd(it):
            try:
                return float(it.get("quantity") or 0)
            except (TypeError, ValueError):
                return 0.0

        for _tipo, _pref, _its in _agrupar_por_prefixo(_ordenados):
            if _tipo == "grupo":
                # UM bloco pro conjunto: total + ambientes. Regra nº1 no
                # agregado: "medido" só se TODOS forem; regra nº4 preservada —
                # cada quantidade continua listada, a prosa é que agrega.
                _un = (_its[0].get("unit") or "").strip()
                _tot = sum(_qtd(x) for x in _its)
                _todos_medidos = all((x.get("confidence") or "") == "confirmado" for x in _its)
                _rot = "medido do CAD" if _todos_medidos else "estimativa — a confirmar"
                _partes = []
                for x in _its:
                    _resto = _re.split(r"\s+—\s+", (x.get("description") or "").strip(), maxsplit=1)
                    _amb = _poda_boilerplate(_resto[1] if len(_resto) > 1 else "").strip() or "item"
                    _q = _qtd(x)
                    _partes.append(f"{_amb} ({_fmt_qty(_q)} {_un})" if _q > 0 else f"{_amb} (a confirmar)")
                texto = (f"{_pref} — {len(_its)} itens, total {_fmt_qty(_tot)} {_un} "
                         f"({_rot}): " + "; ".join(_partes) + ".")
                blocos.append({"tipo": "item", "texto": texto,
                               "origem": "medido" if _todos_medidos else "estimado"})
                continue
            it = _its[0]
            qty = _qtd(it)
            unit = (it.get("unit") or "").strip()
            medido = (it.get("confidence") or "") == "confirmado"
            if qty > 0:
                origem_txt = "medido do CAD" if medido else "estimativa — a confirmar"
                sufixo = f" — {_fmt_qty(qty)} {unit} ({origem_txt})"
            else:
                sufixo = " — quantidade a confirmar"
            texto = _poda_boilerplate((it.get("description") or "").strip()) + sufixo
            obs = (it.get("observations") or "").strip()
            if obs:
                texto += f" Obs.: {obs[:220]}"
            blocos.append({"tipo": "item", "texto": texto,
                           "origem": "medido" if medido else "estimado"})
        blocos.append(_bloco_preencher(LACUNAS_POR_DISCIPLINA.get(disc, LACUNA_PADRAO)))
        secoes.append({"titulo": f"{num}. {disc}", "blocos": blocos})
        num += 1

    secoes.append({"titulo": f"{num}. Considerações finais", "blocos": [
        {"tipo": "texto", "texto": (
            "Havendo divergência entre este memorial e os desenhos do projeto, prevalecem os desenhos. "
            "As quantidades assinaladas como estimativa devem ser conferidas pelo responsável técnico "
            "antes de qualquer uso contratual, bancário ou de aprovação. Este documento é um rascunho "
            "gerado automaticamente a partir do arquivo CAD e não possui validade como memorial "
            "descritivo até ser revisado, complementado e assinado pelo responsável técnico.")},
    ]})
    secoes.append({"titulo": f"{num + 1}. Encerramento e assinaturas", "blocos": [
        {"tipo": "texto", "texto": "Local e data: ____________________________, ____/____/______"},
        {"tipo": "assinatura", "texto": "Proprietário / contratante"},
        {"tipo": "assinatura", "texto": "Responsável técnico — registro e RRT/ART nº [A PREENCHER]"},
    ]})

    return {
        "versao": 1,
        "obra": nome_obra,
        "gerado_em": _agora_brasilia().strftime("%d/%m/%Y"),
        "resumo": {"total": total, "medidos": medidos, "estimados": estimados},
        "secoes": secoes,
    }


# ═══════════════════════════════════════════════════════════════
#  ETAPA 2 — render .docx a partir da estrutura (original ou editada)
# ═══════════════════════════════════════════════════════════════

def estrutura_para_docx(caminho, estrutura: dict):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cabeçalho e rodapé em TODAS as páginas (cerca nº4)
    sec = doc.sections[0]
    sec.top_margin, sec.bottom_margin = Cm(2.2), Cm(2.0)
    sec.left_margin, sec.right_margin = Cm(2.5), Cm(2.5)
    h = sec.header.paragraphs[0]
    h.text = "RASCUNHO — para revisão do responsável técnico. Não substitui o memorial assinado (RRT/ART)."
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = VERMELHO
        run.bold = True
    f = sec.footer.paragraphs[0]
    f.text = (f"Rascunho gerado pelo AI.arq a partir do arquivo CAD do projeto em "
              f"{estrutura.get('gerado_em') or _agora_brasilia().strftime('%d/%m/%Y')}"
              " · quantidades estimadas devem ser confirmadas")
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in f.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = CINZA

    # Capa compacta
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run("MEMORIAL DESCRITIVO")
    rt.bold = True
    rt.font.size = Pt(22)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run((estrutura.get("obra") or "").strip() or "[A PREENCHER: nome da obra]")
    rs.font.size = Pt(14)
    carimbo = doc.add_paragraph()
    carimbo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = carimbo.add_run("RASCUNHO PARA REVISÃO DO RESPONSÁVEL TÉCNICO")
    rc.bold = True
    rc.font.size = Pt(12)
    rc.font.color.rgb = VERMELHO

    n_itens = 0
    for secao in estrutura.get("secoes") or []:
        doc.add_heading((secao.get("titulo") or "").strip() or "—", level=1)
        for bloco in secao.get("blocos") or []:
            tipo = bloco.get("tipo") or "texto"
            texto = (bloco.get("texto") or "").strip()
            if tipo == "tabela":
                linhas = bloco.get("linhas") or []
                if not linhas:
                    continue
                tab = doc.add_table(rows=0, cols=2)
                tab.style = "Table Grid"
                for par in linhas:
                    rotulo = str(par[0]) if len(par) > 0 else ""
                    valor = str(par[1]) if len(par) > 1 else ""
                    cells = tab.add_row().cells
                    cells[0].paragraphs[0].add_run(rotulo).bold = True
                    cells[1].text = valor
                    cells[0].width, cells[1].width = Cm(5.5), Cm(10.5)
            elif tipo == "preencher" or (texto.startswith(MARCA_PREENCHER)):
                p = doc.add_paragraph()
                run = p.add_run(texto)
                run.italic = True
                run.font.color.rgb = LARANJA
                run.font.size = Pt(10)
            elif tipo == "item":
                n_itens += 1
                p = doc.add_paragraph(style="List Bullet")
                # Destaque do sufixo de quantidade quando ainda dá pra achar
                # (edição livre pode ter mudado o texto — aí vai tudo simples).
                pos = texto.rfind(" — ")
                if pos > 0:
                    p.add_run(texto[:pos])
                    rq = p.add_run(texto[pos:])
                    rq.bold = True
                    if bloco.get("origem") == "estimado":
                        rq.font.color.rgb = LARANJA
                else:
                    p.add_run(texto)
            elif tipo == "assinatura":
                doc.add_paragraph()
                doc.add_paragraph("_________________________________________\n" + texto)
            else:
                if texto:
                    doc.add_paragraph(texto)

    doc.save(caminho)
    resumo = dict(estrutura.get("resumo") or {})
    resumo["secoes"] = len(estrutura.get("secoes") or [])
    resumo["itens_no_docx"] = n_itens
    return resumo


def gerar_memorial_docx(caminho, projeto: dict, items: list):
    """Compatibilidade: monta a estrutura padrão e renderiza o .docx."""
    return estrutura_para_docx(caminho, montar_estrutura(projeto, items))


# ═══════════════════════════════════════════════════════════════
#  Render PDF (WeasyPrint — mesmo motor do cronograma em produção)
# ═══════════════════════════════════════════════════════════════

def _esc(s):
    return (str(s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def estrutura_para_pdf_bytes(estrutura: dict) -> bytes:
    """PDF do memorial a partir da estrutura (original ou editada).
    Import do weasyprint fica AQUI dentro (padrão cronograma_render:837)."""
    partes = []
    for secao in estrutura.get("secoes") or []:
        partes.append(f'<h2>{_esc(secao.get("titulo"))}</h2>')
        for b in secao.get("blocos") or []:
            tipo = b.get("tipo") or "texto"
            texto = _esc((b.get("texto") or "").strip())
            if tipo == "tabela":
                linhas = "".join(
                    f'<tr><th>{_esc(l[0] if len(l) > 0 else "")}</th>'
                    f'<td>{_esc(l[1] if len(l) > 1 else "")}</td></tr>'
                    for l in (b.get("linhas") or []))
                if linhas:
                    partes.append(f'<table>{linhas}</table>')
            elif tipo == "preencher" or texto.startswith(_esc(MARCA_PREENCHER)):
                partes.append(f'<p class="preencher">{texto}</p>')
            elif tipo == "item":
                cls = "est" if b.get("origem") == "estimado" else ""
                partes.append(f'<p class="item {cls}">• {texto}</p>')
            elif tipo == "assinatura":
                partes.append(f'<p class="ass">_________________________________________<br>{texto}</p>')
            elif texto:
                partes.append(f'<p>{texto}</p>')

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>
@page {{
  size: A4; margin: 2.4cm 2.0cm;
  @top-center {{ content: "RASCUNHO — para revisão do responsável técnico. Não substitui o memorial assinado (RRT/ART)."; color: #B91C1C; font-size: 7.5pt; font-weight: 600; }}
  @bottom-center {{ content: "Rascunho gerado pelo AI.arq em {_esc(estrutura.get('gerado_em') or _agora_brasilia().strftime('%d/%m/%Y'))} · página " counter(page) " de " counter(pages); color: #6B7280; font-size: 7.5pt; }}
}}
body {{ font-family: Helvetica, Arial, sans-serif; font-size: 10.5pt; color: #1F2937; line-height: 1.45; }}
h1 {{ text-align: center; font-size: 21pt; margin: 0 0 2pt; }}
.obra {{ text-align: center; font-size: 13pt; margin: 0 0 4pt; color: #374151; }}
.carimbo {{ text-align: center; font-size: 11pt; font-weight: 700; color: #B91C1C; margin: 0 0 18pt; }}
h2 {{ font-size: 12pt; margin: 16pt 0 6pt; padding-bottom: 3pt; border-bottom: 2pt solid #7C3AED; page-break-after: avoid; }}
p {{ margin: 4pt 0; }}
.item {{ margin: 2pt 0 2pt 10pt; }}
.item.est {{ color: #92400E; }}
.preencher {{ color: #B45B09; font-style: italic; background: #FFF7ED; border: 1pt dashed #FDBA74; border-radius: 4pt; padding: 5pt 7pt; }}
.ass {{ text-align: center; margin-top: 26pt; }}
table {{ width: 100%; border-collapse: collapse; margin: 6pt 0; }}
th, td {{ border: 0.7pt solid #D1D5DB; padding: 5pt 7pt; font-size: 10pt; text-align: left; }}
th {{ background: #F9FAFB; width: 38%; font-weight: 600; }}
</style></head><body>
<h1>MEMORIAL DESCRITIVO</h1>
<p class="obra">{_esc(estrutura.get("obra") or "")}</p>
<p class="carimbo">RASCUNHO PARA REVISÃO DO RESPONSÁVEL TÉCNICO</p>
{''.join(partes)}
</body></html>"""
    from weasyprint import HTML
    return HTML(string=html).write_pdf()


# ═══════════════════════════════════════════════════════════════
#  Redação por IA — SOMENTE do parágrafo de abertura de cada seção
# ═══════════════════════════════════════════════════════════════
# Pedro (02/08): "não podemos usar a IA pra fazer junto? não ficaria melhor?"
# Fica — mas com coleira. O v1 determinístico escrevia a MESMA frase genérica
# em toda disciplina ("Os serviços de X compreendem os itens..."), o que lê
# como lista, não como memorial. A IA melhora exatamente isso.
#
# O QUE A IA FAZ: reescreve só o parágrafo de ABERTURA de cada disciplina,
# resumindo em prosa técnica o que aquela seção contém — a partir das
# descrições dos itens daquele projeto.
# O QUE ELA NUNCA TOCA: as linhas de item (quantidade, unidade, rótulo
# medido/estimado), os blocos [A PREENCHER], o carimbo de rascunho, as
# tabelas e o encerramento. Ou seja: o número nunca passa pela IA.
#
# 🚫 NÚMERO NENHUM no parágrafo. É a regra que torna a invenção detectável:
# se aparecer dígito, o texto é rejeitado inteiro (validador abaixo). Assim a
# IA não tem como inventar quantidade, bitola, diâmetro ou norma numerada.

PROMPT_REDACAO = """Você escreve memoriais descritivos de obra em português do Brasil.

Receberá a lista REAL de itens levantados de uma disciplina do projeto. Escreva
UM único parágrafo de abertura para essa seção do memorial, entre 25 e 60 palavras.

REGRAS ABSOLUTAS (violar qualquer uma invalida sua resposta):
1. Escreva SOMENTE sobre o que está nos itens. Não acrescente material,
   sistema, marca, norma, NBR, método executivo ou característica que não
   apareça na lista.
2. NÃO escreva NENHUM número, algarismo ou quantidade. Nem por extenso com
   valor (ex.: "doze metros"). As quantidades ficam na lista abaixo do
   parágrafo — não repita nada disso.
3. Não invente cor, dimensão, espessura, traço, fabricante ou desempenho.
4. Não prometa qualidade ("de alto padrão", "premium", "excelente") nem faça
   juízo. Texto técnico, neutro, impessoal, no futuro ("serão executados").
5. Não cite lei, norma, decreto ou sigla técnica que não esteja nos itens.

Responda APENAS com o parágrafo, sem título, sem aspas, sem comentário."""

# Palavras/padrões que denunciam invenção — checados no texto que volta.
_PROIBIDOS_RX = None


def _rx_proibidos():
    global _PROIBIDOS_RX
    if _PROIBIDOS_RX is None:
        _PROIBIDOS_RX = re.compile(
            r"\d"                                   # qualquer algarismo
            r"|\bnbr\b|\babnt\b|\biso\b|\bnr[\s-]?\d*\b"   # normas
            r"|\bconforme\s+norma\b|\blei\b|\bdecreto\b"
            r"|\bpremium\b|\balto\s+padr[aã]o\b|\bexcelente\b|\bqualidade\s+superior\b",
            re.I)
    return _PROIBIDOS_RX


def validar_paragrafo_ia(texto: str) -> tuple:
    """(ok, motivo). Rejeita qualquer parágrafo com número, norma ou juízo de
    valor — e limita tamanho. Na dúvida REJEITA: o texto determinístico é o
    fallback e ele nunca mente."""
    t = (texto or "").strip()
    if not t:
        return False, "vazio"
    if len(t) > 700:
        return False, "longo demais"
    n_palavras = len(t.split())
    if n_palavras < 12 or n_palavras > 90:
        return False, f"tamanho fora da faixa ({n_palavras} palavras)"
    m = _rx_proibidos().search(t)
    if m:
        return False, f"contém termo proibido: {m.group(0)!r}"
    if "\n" in t.strip():
        return False, "mais de um parágrafo"
    return True, ""


def redigir_intros_ia(estrutura: dict, chamar_llm) -> dict:
    """Reescreve o 1º bloco de texto das seções de DISCIPLINA usando IA.

    chamar_llm(system, user) -> str  (injetado pelo main.py; aqui não há
    dependência de SDK, o que mantém este módulo testável sem rede).

    Devolve {"reescritas": n, "puladas": [(secao, motivo)]} e MUTA a estrutura.
    Seção sem item, ou cuja resposta não passa no validador, fica com o texto
    determinístico — nunca quebra, nunca degrada pra pior."""
    fixas = ("apresenta", "dados da obra", "responsabilidade",
             "considerações finais", "encerramento")
    reescritas, puladas = 0, []
    for secao in estrutura.get("secoes") or []:
        titulo = (secao.get("titulo") or "")
        nome = re.sub(r"^\s*\d+\.\s*", "", titulo).strip()
        if any(f in nome.lower() for f in fixas):
            continue
        blocos = secao.get("blocos") or []
        itens = [b for b in blocos if b.get("tipo") == "item"]
        if not itens:
            continue
        alvo = next((b for b in blocos if b.get("tipo") == "texto"), None)
        if not alvo:
            continue
        # Só as DESCRIÇÕES — a quantidade nem entra no prompt (não há o que inventar)
        descricoes = []
        for b in itens[:40]:
            d = (b.get("texto") or "").split(" — ")[0].strip()
            if d:
                descricoes.append(f"- {d[:160]}")
        user = (f"Disciplina: {nome}\n\nItens levantados no projeto:\n"
                + "\n".join(descricoes))
        try:
            resp = chamar_llm(PROMPT_REDACAO, user)
        except Exception as e:
            puladas.append((nome, f"erro na chamada: {e}"))
            continue
        ok, motivo = validar_paragrafo_ia(resp)
        if not ok:
            puladas.append((nome, motivo))
            continue
        alvo["texto"] = resp.strip()
        alvo["ia"] = True
        reescritas += 1
    return {"reescritas": reescritas, "puladas": puladas}


# ═══════════════════════════════════════════════════════════════
#  Validação da estrutura EDITADA (vinda do navegador — não confiar)
# ═══════════════════════════════════════════════════════════════

TIPOS_VALIDOS = {"texto", "preencher", "item", "tabela", "assinatura"}
MAX_SECOES = 80
MAX_BLOCOS_POR_SECAO = 600
MAX_TEXTO = 4000
MAX_LINHAS_TABELA = 40


def validar_estrutura_editada(est) -> dict:
    """Sanitiza a estrutura vinda do cliente. Levanta ValueError se inválida.
    Só os campos conhecidos sobrevivem (nada de HTML/campos extras no banco)."""
    if not isinstance(est, dict):
        raise ValueError("estrutura deve ser objeto")
    secoes = est.get("secoes")
    if not isinstance(secoes, list) or not secoes:
        raise ValueError("estrutura sem seções")
    if len(secoes) > MAX_SECOES:
        raise ValueError("seções demais")

    def _txt(v, cap=MAX_TEXTO):
        return str(v or "")[:cap].strip()

    saida = {
        "versao": 1,
        "obra": _txt(est.get("obra"), 200),
        "gerado_em": _txt(est.get("gerado_em"), 20),
        "resumo": {},
        "secoes": [],
    }
    r = est.get("resumo") or {}
    for k in ("total", "medidos", "estimados"):
        try:
            saida["resumo"][k] = int(r.get(k) or 0)
        except (TypeError, ValueError):
            saida["resumo"][k] = 0
    for s in secoes:
        if not isinstance(s, dict):
            raise ValueError("seção inválida")
        blocos_in = s.get("blocos")
        if not isinstance(blocos_in, list) or len(blocos_in) > MAX_BLOCOS_POR_SECAO:
            raise ValueError("blocos inválidos")
        blocos = []
        for b in blocos_in:
            if not isinstance(b, dict):
                raise ValueError("bloco inválido")
            tipo = b.get("tipo")
            if tipo not in TIPOS_VALIDOS:
                raise ValueError(f"tipo de bloco desconhecido: {tipo}")
            novo = {"tipo": tipo}
            if tipo == "tabela":
                linhas = b.get("linhas")
                if not isinstance(linhas, list) or len(linhas) > MAX_LINHAS_TABELA:
                    raise ValueError("tabela inválida")
                novo["linhas"] = [[_txt(l[0] if len(l) > 0 else "", 200),
                                   _txt(l[1] if len(l) > 1 else "", 500)]
                                  for l in linhas if isinstance(l, (list, tuple))]
            else:
                novo["texto"] = _txt(b.get("texto"))
                if tipo == "item" and b.get("origem") in ("medido", "estimado"):
                    novo["origem"] = b["origem"]
            blocos.append(novo)
        saida["secoes"].append({"titulo": _txt(s.get("titulo"), 200) or "—", "blocos": blocos})
    return saida
